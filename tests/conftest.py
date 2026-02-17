"""Shared fixtures for tests — synthetic documents, no network calls."""

from __future__ import annotations

import textwrap
from pathlib import Path

import pytest

from rag.chunking.schemas import ChunkMetadata
from rag.documents.schemas import DocType, DocumentMetadata

# ---------------------------------------------------------------------------
# Directory fixtures
# ---------------------------------------------------------------------------


@pytest.fixture
def tmp_dir(tmp_path: Path) -> Path:
    """Provide a temporary directory."""
    return tmp_path


# ---------------------------------------------------------------------------
# Synthetic document content
# ---------------------------------------------------------------------------


@pytest.fixture
def sample_txt_content() -> str:
    return textwrap.dedent("""\
        Apple Inc. (AAPL) — Q4 2024 Earnings Summary

        Revenue came in at $94.9 billion, up 6% year-over-year. Services revenue
        hit a new all-time record of $25.0 billion, driven by advertising, App Store,
        and cloud services. Gross margin expanded 120 basis points to 46.2%.

        Management raised guidance for Q1 2025, citing strong iPhone 16 demand
        and continued growth in emerging markets. The company repurchased
        $25 billion of stock during the quarter.

        Risk Factors

        Foreign exchange headwinds remain a concern, with the strong dollar
        reducing international revenue by approximately 3 percentage points.
        Supply chain constraints in advanced chip manufacturing could impact
        product availability during peak holiday season.

        Valuation

        At current levels, AAPL trades at 28x forward earnings, a premium to
        the S&P 500 average of 22x. Our DCF model implies fair value of $195,
        suggesting limited upside from current levels.
    """)


@pytest.fixture
def sample_txt_file(tmp_path: Path, sample_txt_content: str) -> Path:
    p = tmp_path / "apple_report.txt"
    p.write_text(sample_txt_content)
    return p


@pytest.fixture
def sample_pdf_file(tmp_path: Path) -> Path:
    """Create a minimal PDF using fpdf2."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Page 1
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.multi_cell(0, 10, text=(
        "Apple Inc. Form 10-K\n\n"
        "Item 1. Business\n\n"
        "Apple Inc. designs, manufactures, and markets smartphones, "
        "personal computers, tablets, wearables, and accessories worldwide. "
        "The company offers iPhone, Mac, iPad, and wearables."
    ))

    # Page 2
    pdf.add_page()
    pdf.multi_cell(0, 10, text=(
        "Item 1A. Risk Factors\n\n"
        "The company faces risks related to global economic conditions, "
        "competitive pressures, and supply chain disruptions. Foreign "
        "exchange fluctuations can materially impact reported revenue "
        "and earnings in any given quarter."
    ))

    # Page 3
    pdf.add_page()
    pdf.multi_cell(0, 10, text=(
        "Item 7. Management's Discussion and Analysis\n\n"
        "Revenue increased 6% year-over-year to $383 billion. Services "
        "segment grew 14% and now represents 26% of total revenue. "
        "Operating margin expanded to 30.7% from 29.8% in the prior year."
    ))

    p = tmp_path / "apple_10k.pdf"
    pdf.output(str(p))
    return p


@pytest.fixture
def sample_docx_file(tmp_path: Path) -> Path:
    """Create a minimal DOCX file."""
    from docx import Document

    doc = Document()
    doc.add_heading("Equity Research Report", level=1)
    doc.add_paragraph(
        "Apple Inc. (AAPL) remains our top pick in large-cap technology. "
        "The services flywheel continues to accelerate, with recurring "
        "revenue now exceeding $100 billion annualized."
    )
    doc.add_heading("Valuation", level=2)
    doc.add_paragraph(
        "Our price target of $220 is based on a 30x multiple applied to "
        "our CY2025 EPS estimate of $7.35. This implies 15% upside."
    )

    p = tmp_path / "aapl_research.docx"
    doc.save(str(p))
    return p


@pytest.fixture
def sample_xlsx_file(tmp_path: Path) -> Path:
    """Create a minimal Excel file with financial data."""
    import pandas as pd

    revenue_data = {
        "Year": [2022, 2023, 2024],
        "Revenue ($B)": [394.3, 383.3, 395.8],
        "Gross Margin": ["43.3%", "44.1%", "46.2%"],
        "EPS": [6.11, 6.16, 6.97],
    }
    df_revenue = pd.DataFrame(revenue_data)

    segment_data = {
        "Segment": ["iPhone", "Services", "Mac", "iPad", "Wearables"],
        "Revenue ($B)": [200.6, 85.2, 29.4, 28.3, 39.8],
        "YoY Growth": ["2%", "14%", "5%", "-3%", "-1%"],
    }
    df_segments = pd.DataFrame(segment_data)

    p = tmp_path / "aapl_model.xlsx"
    with pd.ExcelWriter(str(p), engine="openpyxl") as writer:
        df_revenue.to_excel(writer, sheet_name="Summary", index=False)
        df_segments.to_excel(writer, sheet_name="Segments", index=False)

    return p


@pytest.fixture
def sec_filing_text() -> str:
    """Multi-page SEC filing text with ITEM boundaries."""
    return (
        "PART I\n\n"
        "Item 1. Business\n\n"
        "Apple Inc. designs smartphones and computers. "
        "The company operates globally with significant presence "
        "in North America, Europe, and Greater China. "* 20 + "\n\n"
        "Item 1A. Risk Factors\n\n"
        "Global economic conditions affect demand for consumer electronics. "
        "Foreign exchange fluctuations impact international revenue. "
        "Supply chain disruptions can affect product availability. " * 15 + "\n\n"
        "Item 7. Management's Discussion and Analysis\n\n"
        "Revenue for fiscal 2024 was $395.8 billion, an increase of 3.3% "
        "from $383.3 billion in fiscal 2023. Services revenue grew 14% "
        "year-over-year to $85.2 billion. " * 25 + "\n\n"
        "Item 8. Financial Statements\n\n"
        "Consolidated Balance Sheet as of September 28, 2024. "
        "Total assets: $352.6 billion. Total liabilities: $290.4 billion."
    )


@pytest.fixture
def earnings_transcript_text() -> str:
    """Sample earnings transcript with prepared remarks and Q&A."""
    return textwrap.dedent("""\
        Apple Inc. Q4 2024 Earnings Call

        Tim Cook -- Chief Executive Officer

        Good afternoon everyone. Thank you for joining us. We are pleased to
        report another record quarter with revenue of $94.9 billion. Our
        services business reached an all-time high and our installed base
        continues to grow across all product categories.

        Luca Maestri -- Chief Financial Officer

        Thank you Tim. Total revenue was $94.9 billion, up 6% year-over-year.
        Products revenue was $69.9 billion and services revenue reached a
        new all-time record of $25.0 billion. Gross margin was 46.2%.

        Question-and-Answer Session

        Amit Daryanani -- Evercore ISI

        Congratulations on the quarter. Can you talk about the trajectory
        of services revenue and how you see growth sustaining into fiscal 2025?

        Luca Maestri -- Chief Financial Officer

        Sure Amit. We expect services revenue to maintain double-digit growth.
        The key drivers are advertising, App Store, and cloud services. We now
        have over 1 billion paid subscriptions across our platforms.
    """)


@pytest.fixture
def sample_document_metadata() -> DocumentMetadata:
    return DocumentMetadata(
        doc_type=DocType.SEC_FILING,
        ticker="AAPL",
        filing_date="2024-10-30",
        filing_type="10-K",
    )


@pytest.fixture
def sample_chunk_metadata() -> ChunkMetadata:
    return ChunkMetadata(
        doc_type=DocType.SEC_FILING,
        ticker="AAPL",
        filing_date="2024-10-30",
        section_name="Risk Factors",
        item_number="1A",
    )
